#!/usr/bin/env python3

"""
This is a python file for the OpenCAP

Usage: python3 identified_variants_to_annotation.py <somatic variants file>

Arguments:
    <somatic variants file> = TSV (tab separated values) of putative somatic variants with 5 columns [chromosome, start, stop, reference, variant].
    
"""

from docx import Document
from docx.shared import Inches
import pandas as pd
import datetime
import solvebio
import myvariant
import utils

# Pull in variant file
somatic_variants = pd.read_csv(sys.argv[1], sep='\t')
sample_name = sys.argv[2]

def myvariantinfo(chrom, start, ref, var):
    variant = myvariant.format_hgvs(chrom, start, ref, var)
    directory = mv.getvariant(variant)
    if directory:
        gene = directory['snpeff']['ann'][0]['gene_id']
        varinat_type = directory['cadd']['consdetail']
        consequence = directory['cadd']['consequence']
        ENST = directory['cadd']['gene']['feature_id']
        ENSG = directory['cadd']['gene']['gene_id']
        coding_change = directory['snpeff']['ann'][-1]['hgvs_c']
        if directory['snpeff']['ann'][-1]['hgvs_p']:
            protein_change = directory['snpeff']['ann'][-1]['hgvs_p'] 
        pop_freq = directory['gnomad_exome']['af']['af']
        information = [gene, varinat_type, consequence, ENSG, ENST, coding_change, protein_change, pop_freq]
        return information
    else:
        return 'None'

processed = 0
clinical = 0

clinical_info = []
for i,row in somatic_variants.iterrows():
    chrom = row['Chromosome']
    start = int(row['Start'])
    ref = row['Ref']
    var = row['Var']
    result = myvariantinfo(chrom, start, ref, var)
    if result == 'None':
        processed +=1
    if result != 'None':
        clinical +=1
        clinical_info.append(result)

clinical_info = pd.DataFrame(clinical_info)
clinical_info.columns=['','','','','']

# def build_document(sample_name, input_list, CIViC_annotations):
currentDT = datetime.datetime.now()

#Initiate document
document = Document()

document.add_picture('Extra/report_header.png', width=Inches(6))
document.add_heading('SOMATIC VARIANT ANNOTATION', 0)

#Add Information about Sample and Processing
p = document.add_paragraph()
p.add_run('Sample Name: ').bold = True
p.add_run(str(sample_name) + '\n')
p.add_run('Date ').bold = True
p.add_run(str(currentDT.strftime("%a, %b %d, %Y")) + '\n')
p.add_run('Time Processed: ').bold = True
p.add_run(str(currentDT.strftime("%I:%M:%S %p")) + '\n')
p.add_run('Variants Processed: ').bold = True
p.add_run(str(processed) + '\n')
p.add_run('Clinical Annotations: ').bold = True
p.add_run(str(clinical) + '\n')

#Add Direct CIViC Annotations
document.add_heading('Direct CIViC Annotations' + '\n', level=1) 

table = document.add_table(rows=1, cols=5, style = 'Table Grid')
hdr_cells = table.rows[0].cells
run = hdr_cells[0].paragraphs[0].add_run('Gene')
run.bold = True
run = hdr_cells[1].paragraphs[0].add_run('Variant')
run.bold = True
run = hdr_cells[2].paragraphs[0].add_run('Description')
run.bold = True
run = hdr_cells[3].paragraphs[0].add_run('CIViC EID')
run.bold = True
run = hdr_cells[4].paragraphs[0].add_run('PubMedID')
run.bold = True

for i, row in clinical_info.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Gene']
    row_cells[1].text = row['Variant']
    row_cells[2].text = row['Description']
    row_cells[3].text = row['CIViC Link']
    row_cells[4].text = row['PubMedID Link']

#Add Indirect CIVIC titles
document.add_heading('Indirect CIViC Annotations' + '\n', level=1) 


    
p = document.add_paragraph('\n' + 'OpenCAP is intended for research use only and clinical applications of subsequent panels designed using the SOP would require further panel validation.')

document.save('demo.docx')



# Pull in CIViC API for all variants
#variants_DNA = requests.get('https://civic.genome.wustl.edu/api/panels/DNA-based/qualifying_variants?minimum_score=0').json()['records']