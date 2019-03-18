#!/usr/bin/env python3

"""
This is a python file for the OpenCAP

Usage: python3 identified_variants_to_annotation.py <somatic variants file> <sample name>

Arguments:
    <somatic variants file> = TSV (tab separated values) of putative somatic variants with 5 columns [chromosome, start, stop, reference, variant].
    
    <sample name> = label the sample that is being annotated
    
"""
import sys
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd
import datetime
import myvariant

# Pull in variant file
somatic_variants = pd.read_csv(sys.argv[1], sep='\t')
sample_name = sys.argv[2]

#Crate myvarinat info pull
mv = myvariant.MyVariantInfo()

# Create function to pull evidence statements from CIViC Directory derived from myvariant info
def get_evidence_statements(civic_directory):
    evidence_statements = []
    assertion_statements = []
    
    sample_evidence_count = {}
    sample_evidence_count = [0,0,0,0]
    
    gene = str(civic_directory['entrez_name'])
    variant = str(civic_directory['name'])
    
    for evidence in civic_directory['evidence_items']:
        if evidence['evidence_level'] != 'D' and evidence['evidence_level'] != 'E' and evidence['status'] == 'accepted':
            if 'clinical_significance' not in evidence:
                initial = str(gene) +  ' ' +str(variant) +  ' ' +evidence['evidence_direction']+ ' ' + 'outcome'
            else:    
                initial = str(gene) +  ' ' +str(variant) +  ' ' +evidence['evidence_direction']+ ' ' + evidence['clinical_significance']
            CIViC_EID =  evidence['name']
            PubmedID = str(evidence['source']['pubmed'])

            #PREDICTIVE
            if evidence['evidence_type'] == 'Predictive':
                if 'drug_interaction_type' not in evidence:
                    evidence_statements.append([initial + ' to ' + evidence['drugs'][0]['name'] + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID])

                elif evidence['drug_interaction_type'] == 'Combination':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + 'combination of ' + ', '.join(drug_list[:-1]) + ' and ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID])


                elif evidence['drug_interaction_type'] == 'Substitutes':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + ', '.join(drug_list[:-1]) + ' or ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID])

                elif evidence['drug_interaction_type'] == 'Substitutes':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + ', '.join(drug_list[:-1]) + ' or ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID])

                elif evidence['drug_interaction_type'] == 'Sequential':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + 'sequence of ' + ', '.join(drug_list[:-1]) + ' and ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID])


            #CREATE PROGNOSTIC EVIDENCE STATEMENT
            if evidence['evidence_type'] == 'Prognostic':
                evidence_statements.append([initial + ' for patients with ' + evidence['disease']['name'], CIViC_EID, PubmedID])


            #CREATE DIAGNOSTIC EVIDENCE STATEMENT
            if evidence['evidence_type'] == 'Diagnostic':
                evidence_statements.append([initial + ' for patients with ' + evidence['disease']['name'], CIViC_EID, PubmedID])


            #CREATE PREDISPOSING EVIDENCE STATEMENT
            if evidence['evidence_type'] == 'Predisposing':
                evidence_statements.append([initial + ' Predisposition For Cancer ' + ' for patients with ' + evidence['disease']['name'], CIViC_EID, PubmedID])


            if evidence['evidence_type'] == 'Predictive':
                sample_evidence_count[0] += 1
            if evidence['evidence_type'] == 'Prognostic':
                sample_evidence_count[1] += 1
            if evidence['evidence_type'] == 'Diagnostic':
                sample_evidence_count[2] += 1
            if evidence['evidence_type'] == 'Predisposing':
                sample_evidence_count[3] += 1
    
    final_evidence = {}
    for item in evidence_statements:
        if item[0] not in final_evidence:
            final_evidence[item[0]] = [[item[1]],[item[2]]]
        else:
            final_evidence[item[0]][0].append(item[1])
            final_evidence[item[0]][1].append(item[2])
            
    return final_evidence, sample_evidence_count


# INITIATE DOCUMENT
currentDT = datetime.datetime.now()

document = Document()
document.add_picture('../Extra/report_header.png', width=Inches(6))
document.add_heading('SOMATIC VARIANT ANNOTATION', 0)

# ADD SAMPLE NAME, DATE, AND TIME
p = document.add_paragraph()
p.add_run('Sample Name' + '\t'+ '\t'+ '\t').bold = True
p.add_run(str(sample_name) + '\n')
p.add_run('Date ' + '\t'+ '\t'+ '\t'+ '\t').bold = True
p.add_run(str(currentDT.strftime("%a, %b %d, %Y")) + '\n')
p.add_run('Time Processed ' + '\t'+ '\t').bold = True
p.add_run(str(currentDT.strftime("%I:%M:%S %p")))

# Start running list for variant counts
processed = 0
clinical_count = 0

# Iterate through variant list and pull information
for i,row in somatic_variants.iterrows():
    processed +=1
    
    chrom = row['Chromosome']
    start = int(row['Start'])
    ref = row['Ref']
    var = row['Var']
    
    variant = myvariant.format_hgvs(chrom, start, ref, var)
    directory = mv.getvariant(variant)

    # Pull CIViC Data
    if directory:
        if 'civic' in directory:
            variant_descriptions = []
            assertions = []
            # Add count to clinical
            clinical_count +=1

            # Pull general information for variant
            gene = directory['civic']['entrez_name']
            ENST = directory['cadd']['gene']['feature_id']
            ENSG = directory['cadd']['gene']['gene_id']
            protein_change = directory['civic']['name']
            
            # Pull assertion information for variant
            if 'assertions' in directory['civic'].keys():
                for item in directory['civic']['assertions']:
                    assertions.append(item['description']) 
            
            # Pull evidence item information for variant
            if directory['civic']['description']:
                variant_descriptions.append(directory['civic']['description']) 

            evidence_items, sample_evidence_count = get_evidence_statements(directory['civic'])

            # ADD INFORMATION TO DOCUMENT ABOUT VARIANT
            # general variant information
            document.add_heading('Clinical Variant #' + str(clinical_count) , 1)
            
            p = document.add_paragraph()
            p.add_run('Gene Name' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(gene) + '\n')
            p.add_run('Protein Change' + '\t'+ '\t').bold = True
            p.add_run(str(protein_change) + '\n')
            p.add_run('Coordinates' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(variant) + '\n')
            p.add_run('ENST ID' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(ENST) + '\n')
            p.add_run('ENSG ID' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(ENSG))

            
            document.add_heading('Variant Description: ', 3)
            if not variant_descriptions:
                p = document.add_paragraph('N/A')
            
            else:
                for i,item in enumerate(variant_descriptions):
                    p = document.add_paragraph()
                    p.add_run(str(item))
            
            
            # Assertions
            document.add_heading('Associated Assertions:', 3)
            if not assertions:
                p = document.add_paragraph('N/A')           
            else:
                for i,item in enumerate(assertions):
                    p = document.add_paragraph(style='List Number')
                    p.add_run(str(item) + '\n')
                
            # Evidence Statements
            document.add_heading('Associated Evidence Items:', 3)
            
            table = document.add_table(rows=1, cols=3, style = 'Table Grid')
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            run = hdr_cells[0].paragraphs[0].add_run('Description')
            run.bold = True
            run = hdr_cells[1].paragraphs[0].add_run('CIViC EID')
            run.bold = True
            run = hdr_cells[2].paragraphs[0].add_run('PubMedID')
            run.bold = True
                
            for k,v in evidence_items.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(k)
                row_cells[1].text = str(", ".join(v[0]))
                row_cells[2].text = str(", ".join(v[1]))
                
            for cell in table.columns[0].cells:
                cell.width = Inches(4.5)
            for cell in table.columns[1].cells:
                cell.width = Inches(1.5)
            for cell in table.columns[2].cells:
                cell.width = Inches(1.5)

document.add_heading('Processing information', 1)

# ADD FINAL INFORMATION
p = document.add_paragraph()
# p.add_run('Sample Name: ').bold = True
p.add_run('Variants Processed: ').bold = True
p.add_run(str(processed) + '\n')
p.add_run('Clinical Annotations: ').bold = True
p.add_run(str(clinical_count) + '\n')

# ADD DISCLAIMER    
p = document.add_paragraph('OpenCAP is intended for research use only and clinical applications of subsequent panels designed using the SOP would require further panel validation.').italics = True

sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)
    section.bottom_margin = Inches(.5)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)

# SAVe DOCUMENTS
document.save(sample_name + '_OpenCAP_report.docx')   