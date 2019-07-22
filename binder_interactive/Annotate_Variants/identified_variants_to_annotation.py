#!/usr/bin/env python3

"""
This is a python file for the OpenCAP

Usage: python3 identified_variants_to_annotation.py <somatic variants file> <sample name>

Arguments:
    <somatic variants file> = TSV (tab separated values) of putative somatic variants with 5 columns [chromosome, start, stop, reference, variant].
    
    <sample name> = label the sample that is being annotated
    
"""
import sys
from docx import Document, oxml, opc
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import pandas as pd
import datetime
import myvariant
import requests
import warnings
warnings.filterwarnings('ignore')


# Pull in variant file
somatic_variants = pd.read_csv(sys.argv[1], sep='\t')
sample_name = sys.argv[2]

#Crate myvarinat info pull
mv = myvariant.MyVariantInfo()

# Create function to add hyperlinks for evidence IDs
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = oxml.shared.OxmlElement('w:r')
    rPr = oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

# Create function to pull evidence statements from CIViC Directory derived from myvariant info
def get_evidence_statements(civic_directory):
    evidence_statements = []
    assertion_statements = []
    
    sample_evidence_count = {}
    sample_evidence_count = [0,0,0,0]
    
    gene = str(civic_directory['entrez_name'])
    variant = str(civic_directory['name'])
    
    for evidence in civic_directory['evidence_items']:
        if evidence['evidence_level'] != 'D' and evidence['evidence_level'] != 'E' and evidence['status'] == 'accepted':
            if 'clinical_significance' not in evidence:
                initial = str(gene) +  ' ' +str(variant) +  ' ' +evidence['evidence_direction']+ ' ' + 'outcome'
            else:    
                initial = str(gene) +  ' ' +str(variant) +  ' ' +evidence['evidence_direction']+ ' ' + evidence['clinical_significance']
            CIViC_EID =  evidence['name']
            PubmedID = str(evidence['source']['citation_id'])
            Source = str(evidence['source']['citation'])

            #PREDICTIVE
            if evidence['evidence_type'] == 'Predictive':
                if 'drug_interaction_type' not in evidence:
                    evidence_statements.append([initial + ' to ' + evidence['drugs'][0]['name'] + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID, Source])

                elif evidence['drug_interaction_type'] == 'Combination':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + 'combination of ' + ', '.join(drug_list[:-1]) + ' and ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID, Source])


                elif evidence['drug_interaction_type'] == 'Substitutes':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + ', '.join(drug_list[:-1]) + ' or ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID, Source])

                elif evidence['drug_interaction_type'] == 'Substitutes':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + ', '.join(drug_list[:-1]) + ' or ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID, Source])

                elif evidence['drug_interaction_type'] == 'Sequential':
                        drug_list = []
                        for drug in evidence['drugs']:
                            drug_list.append(drug['name'])
                        evidence_statements.append([initial + ' to ' + 'sequence of ' + ', '.join(drug_list[:-1]) + ' and ' + str(drug_list[-1]) + ' for patients with ' + evidence['disease'][ 'name'], CIViC_EID, PubmedID, Source])


            #CREATE PROGNOSTIC EVIDENCE STATEMENT
            if evidence['evidence_type'] == 'Prognostic':
                evidence_statements.append([initial + ' for patients with ' + evidence['disease']['name'], CIViC_EID, PubmedID, Source])


            #CREATE DIAGNOSTIC EVIDENCE STATEMENT
            if evidence['evidence_type'] == 'Diagnostic':
                evidence_statements.append([initial + ' for patients with ' + evidence['disease']['name'], CIViC_EID, PubmedID, Source])


            #CREATE PREDISPOSING EVIDENCE STATEMENT
            if evidence['evidence_type'] == 'Predisposing':
                evidence_statements.append([initial + ' Predisposition For Cancer ' + ' for patients with ' + evidence['disease']['name'], CIViC_EID, PubmedID, Source])


            if evidence['evidence_type'] == 'Predictive':
                sample_evidence_count[0] += 1
            if evidence['evidence_type'] == 'Prognostic':
                sample_evidence_count[1] += 1
            if evidence['evidence_type'] == 'Diagnostic':
                sample_evidence_count[2] += 1
            if evidence['evidence_type'] == 'Predisposing':
                sample_evidence_count[3] += 1
    
    final_evidence = {}
    for item in evidence_statements:
        if item[0] not in final_evidence:
            final_evidence[item[0]] = [[item[1]], [item[2]], [item[3]]]
        else:
            final_evidence[item[0]][0].append(item[1])
            final_evidence[item[0]][1].append(item[2])
            final_evidence[item[0]][2].append(item[3])
            
    return final_evidence, sample_evidence_count


# INITIATE DOCUMENT
currentDT = datetime.datetime.now()

document = Document()
document.add_picture('../Extra/report_header.png', width=Inches(6))
document.add_heading('SOMATIC VARIANT ANNOTATION', 0)

# ADD SAMPLE NAME, DATE, AND TIME
p = document.add_paragraph()
p.add_run('Sample Name' + '\t'+ '\t'+ '\t').bold = True
p.add_run(str(sample_name) + '\n')
p.add_run('Date ' + '\t'+ '\t'+ '\t'+ '\t').bold = True
p.add_run(str(currentDT.strftime("%a, %b %d, %Y")) + '\n')
p.add_run('Time Processed ' + '\t'+ '\t').bold = True
p.add_run(str(currentDT.strftime("%I:%M:%S %p")))

# Start running list for variant counts
processed = 0
clinical_count = 0

# Iterate through variant list and pull information
for i,row in somatic_variants.iterrows():
    processed +=1
    
    chrom = row['Chromosome']
    start = int(row['Start'])
    ref = row['Ref']
    var = row['Var']
    
    variant = myvariant.format_hgvs(chrom, start, ref, var)
    directory = mv.getvariant(variant)

    # Pull CIViC Data
    if directory:
        if 'civic' in directory:
            variant_descriptions = []
            assertions = []
            # Add count to clinical
            clinical_count +=1

            # Pull general information for variant
            gene = directory['civic']['entrez_name']
            ENST = directory['civic']['coordinates']['representative_transcript']
            protein_change = directory['civic']['name']
            variant_id = directory['civic']['variant_id']
            
            civic_directory = requests.get('https://civicdb.org/api/variants/' + str(variant_id)).json()
            
            # Pull assertion information for variant
            if 'assertions' in directory['civic'].keys():
                for item in directory['civic']['assertions']:
                    assertions.append(item['description']) 
            
            # Pull variant description for variant
            if 'description' in directory['civic']:
                variant_descriptions.append(directory['civic']['description']) 
            
            evidence_items, sample_evidence_count = get_evidence_statements(civic_directory)

            # ADD INFORMATION TO DOCUMENT ABOUT VARIANT
            # general variant information
            document.add_heading('Clinical Variant #' + str(clinical_count) , 1)
            
            p = document.add_paragraph()
            p.add_run('Gene Name' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(gene) + '\n')
            p.add_run('Protein Change' + '\t'+ '\t').bold = True
            p.add_run(str(protein_change) + '\n')
            p.add_run('Coordinates' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(variant) + '\n')
            p.add_run('ENST ID' + '\t'+ '\t'+ '\t').bold = True
            p.add_run(str(ENST) + '\n')
            #p.add_run('ENSG ID' + '\t'+ '\t'+ '\t').bold = True
            #p.add_run(str(ENSG))

            # External Databases
            document.add_heading('External Databases:', 3)
            if 'clinvar' in directory:
                p = document.add_paragraph()
                p.add_run('ClinVar Allele ID: ').bold = True
                variant_id = str(directory['clinvar']['variant_id'])
                add_hyperlink(p, variant_id, f'https://www.ncbi.nlm.nih.gov/clinvar/variation/{variant_id}/')
                p.add_run('\n')
                
                p.add_run('dbSNP ID: ').bold = True
                rsid = directory['clinvar']['rsid']
                add_hyperlink(p, rsid, f'https://www.ncbi.nlm.nih.gov/snp/{rsid}/')
                p.add_run('\n')
                
                if 'cosmic' in directory:
                    p.add_run('COSMIC ID: ').bold = True
                    cosmic_id = directory['cosmic']['cosmic_id'].strip('COSM')
                    add_hyperlink(p, cosmic_id, f'https://cancer.sanger.ac.uk/cosmic/mutation/overview?id={cosmic_id}')
                    p.add_run('\n')
                
            else:
                p = document.add_paragraph('N/A')
        
        
            document.add_heading('CIViC Variant Description: ', 3)
            if not variant_descriptions:
                p = document.add_paragraph('N/A' + '\n')
            
            else:
                for i,item in enumerate(variant_descriptions):
                    p = document.add_paragraph()
                    p.add_run(str(item))
            
            
            # Assertions
            document.add_heading('Associated CIViC Assertions:', 3)
            if not assertions:
                p = document.add_paragraph('N/A' + '\n')           
            else:
                for i,item in enumerate(assertions):
                    p = document.add_paragraph(style='List Number')
                    p.add_run(str(item) + '\n')
               

            
            #Evidence Statements
            document.add_heading('Associated CIViC Evidence Items:', 3)
            if len(evidence_items) == 0:
                p = document.add_paragraph('N/A') 

            else:
                for k,v in evidence_items.items():
                    p = document.add_paragraph()
                    p.add_run('Description: ').bold = True
                    p.add_run(str(k) + '\n')

                    p.add_run('\n' + '\t')
                    
                    
                    CIViC_eid = p.add_run('CIViC ID(s)')
                    CIViC_eid.bold = True
                    CIViC_eid.underline = True
                    p.add_run('\t' + '\t')
                    Citation = p.add_run('Citation(s)')
                    Citation.bold = True
                    Citation.underline = True
                    p.add_run('\n' + '\t')
                    
                    for item in range(len(v[0])):

                        EID = v[0][item].strip("'")

                        add_hyperlink(p, EID, f'https://civicdb.org/links?idtype=evidence&id={EID}')
                        p.add_run('\t' + '\t')

                        pubmed_id = v[1][item].strip("'")
                        source = v[2][item]

                        add_hyperlink(p, source, f'https://www.ncbi.nlm.nih.gov/pubmed/{pubmed_id}')
                        p.add_run('\n'+ '\t')
                    

document.add_heading('Processing information', 1)

# ADD FINAL INFORMATION
p = document.add_paragraph()
# p.add_run('Sample Name: ').bold = True
p.add_run('Variants Processed: ').bold = True
p.add_run(str(processed) + '\n')
p.add_run('Clinical Annotations: ').bold = True
p.add_run(str(clinical_count) + '\n')

# ADD DISCLAIMER  
document.add_heading('Disclaimer', 1)
p = document.add_paragraph('OpenCAP is intended for research use only and clinical applications of subsequent panels designed using the standard operating procedure would require further panel validation. This Report was generated based on an input list of variants. Recommendations for panel development, sequencing approach, variant identification, and variant refinement can be found at www.OpenCAP.org. These methods require the use of reagents, protocols, instruments, software, databases, and other items. A defect or malfunction in any such materials may compromise the accuracy of the report. The data presented in this report is representative of the databases that are queried. Therefore, presented information is limited to the curation within said databases. Additionally, data within and between database might be inaccurate or inconsistent. Specifically, some of the information within or between databases might be conflicting and / or unreliable. Therefore, it is important that the report be interpreted and considered within the clinical context.')
p.italics = True

sections = document.sections
for section in sections:
    section.top_margin = Inches(.5)
    section.bottom_margin = Inches(.5)
    section.left_margin = Inches(.5)
    section.right_margin = Inches(.5)

# SAVe DOCUMENTS
document.save(sample_name + '_OpenCAP_report.docx') 

print('Variant annotation has been successfully completed!')  
